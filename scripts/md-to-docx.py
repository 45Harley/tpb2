"""Convert the CSPS article draft from markdown to .docx"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Read the markdown
with open(r'c:\tpb2\docs\talk-csps-article-draft.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_styled_paragraph(doc, text, is_italic=False, is_bold=False, alignment=None):
    """Add a paragraph with inline formatting for *italic* and **bold** markers."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment

    # Split on italic/bold markers and apply formatting
    # Handle **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)

    if is_italic:
        for run in p.runs:
            run.italic = True
    if is_bold:
        for run in p.runs:
            run.bold = True

    return p

def process_text(text):
    """Clean markdown artifacts from text."""
    # Remove em dashes that are markdown artifacts (keep real ones)
    text = text.replace(' — ', ' \u2014 ')  # proper em dash
    return text

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Title (# heading)
    if line.startswith('# ') and i < 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Georgia'
        p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    # Subtitle line (italic, centered)
    if line.strip().startswith('*A draft') or line.strip().startswith('*The People'):
        text = line.strip().strip('*')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(18)
        i += 1
        continue

    # Section headings (## heading)
    if line.startswith('## '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Georgia'
        i += 1
        continue

    # Block quotes (>)
    if line.startswith('> '):
        quote_lines = []
        while i < len(lines) and (lines[i].rstrip('\n').startswith('> ') or lines[i].rstrip('\n') == '>'):
            ql = lines[i].rstrip('\n')
            if ql == '>':
                quote_lines.append('')
            else:
                quote_lines.append(ql[2:])
            i += 1
        quote_text = process_text(' '.join(l for l in quote_lines if l))
        # Remove any markdown quote attribution
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Handle italic markers in quote
        parts = re.split(r'(\*.*?\*)', quote_text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.italic = True
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
        continue

    # Numbered list items
    if re.match(r'^\d+\. ', line):
        text = process_text(re.sub(r'^\d+\. ', '', line))
        add_styled_paragraph(doc, text)
        p = doc.paragraphs[-1]
        p.paragraph_format.left_indent = Inches(0.5)
        # Prepend number
        num = re.match(r'^(\d+)\. ', line).group(1)
        first_run = p.runs[0]
        first_run.text = num + '. ' + first_run.text
        i += 1
        continue

    # Empty lines
    if line.strip() == '':
        i += 1
        continue

    # Regular paragraphs
    text = process_text(line)

    # Check if this is the closing italic line
    if text.strip().startswith('*') and text.strip().endswith('*') and not text.strip().startswith('**'):
        inner = text.strip()[1:-1]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        run = p.add_run(inner)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        i += 1
        continue

    add_styled_paragraph(doc, text)
    i += 1

output_path = r'c:\dev git\The Rays Converge - CSPS Draft.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
